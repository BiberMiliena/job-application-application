import re # regular expressions library
import string # string library
import wx  # type: ignore # form toolkit
import docx # type: ignore # main work with documents

# styled text for regular labels
def TextRegular(text):
    font = text.GetFont()
    font.PointSize += 3
    text.SetFont(font)

# clears inputted strings from whitespace characters at the beginning and at the end
def CheckForWhitespaces(text):
    beginning = 0
    while (re.search(text[beginning], string.whitespace) != None):
        beginning += 1
    end = len(text) - 1
    while (re.search(text[end], string.whitespace) != None):
        end -= 1
    return text[beginning:(end+1)]

# replaces an instance of "Germany" with "Deutschland"
def Deutschland(text):
    index = text.find("Germany")
    if (index == -1):
        return text
    else:
        return text[0:index] + "Deutschland"

# filling and saving the .docx file
def Save(self):
    document = docx.Document('Template.docx') # template document

    document.paragraphs[3].add_run(CheckForWhitespaces(self.companyName.GetLineText(0)))
    parser = self.companyAddress.GetLineText(0).index(',')
    document.paragraphs[4].add_run(CheckForWhitespaces(self.companyAddress.GetLineText(0)[0:parser]))
    document.paragraphs[5].add_run(CheckForWhitespaces(Deutschland(self.companyAddress.GetLineText(0)[(parser+2):])))

    document.paragraphs[8].add_run(CheckForWhitespaces(self.position.GetLineText(0)) + ' bewerben.')

    if (self.hours.IsChecked() == True):
        document.paragraphs[8].add_run(' Als Student kann ich nur circa 10 Stunden pro Woche arbeiten. Ich hoffe, dass solche Teilzeitbeschäftigungen möglich sind.')

    document.save('Miliena Biber Anschreiben.docx')



# main panel class
class MainPanel(wx.Panel):
    def __init__(self, parent):
        wx.Panel.__init__(self, parent)

        # setting up sizers for object positioning on the form
        mainSizer = wx.BoxSizer(wx.VERTICAL)
        horizontalSizer = wx.BoxSizer(wx.HORIZONTAL)
        gridSizer = wx.GridBagSizer(hgap=5, vgap=5)

        self.label1 = wx.StaticText(self, label="Company name:")
        TextRegular(self.label1)
        gridSizer.Add(self.label1, pos=(0,0))
        self.companyName = wx.TextCtrl(self, size=(500,25))
        gridSizer.Add(self.companyName, pos=(0,1))

        self.label2 = wx.StaticText(self, label="Company address:")
        TextRegular(self.label2)
        gridSizer.Add(self.label2, pos=(1,0))
        self.companyAddress = wx.TextCtrl(self, size=(500,25))
        gridSizer.Add(self.companyAddress, pos=(1,1))

        self.label3 = wx.StaticText(self, label="Position:")
        TextRegular(self.label3)
        gridSizer.Add(self.label3, pos=(2,0))
        self.position = wx.TextCtrl(self, size=(500,25))
        gridSizer.Add(self.position, pos=(2,1))

        self.hours = wx.CheckBox(self, label="10 hours/week")
        TextRegular(self.hours)
        gridSizer.Add(self.hours, pos=(3,0))

        # save .docx file button
        self.button = wx.Button(self, size=(200,50), label="Save")
        self.Bind(wx.EVT_BUTTON, self.OnClick, self.button)
        TextRegular(self.button)

        # combining and finalising sizers
        horizontalSizer.Add(gridSizer, 0, wx.ALL, 5)
        mainSizer.Add(horizontalSizer, 0, wx.ALL, 5)
        mainSizer.Add(self.button, 0, wx.CENTER)
        self.SetSizerAndFit(mainSizer)

    def OnClick(self, event):
        Save(self)



# starting the form
if __name__ == '__main__':

    app = wx.App()
    frame = wx.Frame(None, title='Job Application Application')
    panel = MainPanel(frame)
    frame.Fit()
    frame.Show()
    app.MainLoop()